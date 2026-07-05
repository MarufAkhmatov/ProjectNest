"""Document text extraction for Temur's knowledge base.

Reads whatever the user drops into `knowladgebasefromdocs/` (bank structure,
project/committee/PMO-department regulations, new-product specs) and returns
plain text so rag.py can chunk + embed it.

Supported: PDF (text layer, else OCR), DOCX, images (OCR), XLSX, TXT/MD/CSV/HTML.
OCR uses Tesseract (rus + uzb + uzb_cyrl + eng). Extracted text is cached on
disk keyed by (file hash + extractor version) so re-embedding never re-OCRs an
unchanged file — OCR is the slow part.

Everything degrades gracefully: a missing library or Tesseract just skips that
file instead of crashing the index build.
"""
from __future__ import annotations
import os
import re
import json
import shutil
import hashlib
from pathlib import Path

from . import config

# Bump when extraction logic changes so the text cache invalidates.
EXTRACT_VERSION = "v1"

# The folder the user fills with source documents (their exact spelling).
DOCS_DIR = config.ROOT / "knowladgebasefromdocs"
# Back-compat / alternate spellings we also accept if present.
_ALT_DIRS = [config.ROOT / "knowledgebasefromdocs", config.ROOT / "knowladgebasefordocs"]

_CACHE_DIR = config.TEMP / "docs_cache"
_TEXT_EXT = {".txt", ".md", ".csv", ".tsv", ".log"}
_HTML_EXT = {".html", ".htm"}
_IMG_EXT = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}
_OCR_LANGS = "rus+uzb+uzb_cyrl+eng"


def docs_dirs() -> list[Path]:
    return [d for d in ([DOCS_DIR] + _ALT_DIRS) if d.exists()]


def _tesseract_exe() -> str | None:
    exe = shutil.which("tesseract")
    if exe:
        return exe
    for cand in (r"C:\Program Files\Tesseract-OCR\tesseract.exe",
                 r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"):
        if os.path.exists(cand):
            return cand
    return None


def _ocr_image(img) -> str:
    try:
        import pytesseract
        exe = _tesseract_exe()
        if exe:
            pytesseract.pytesseract.tesseract_cmd = exe
        elif not shutil.which("tesseract"):
            return ""
        return pytesseract.image_to_string(img, lang=_OCR_LANGS) or ""
    except Exception:
        return ""


def _from_pdf(path: Path) -> str:
    """Text layer first; pages with little/no text are rendered and OCR'd."""
    try:
        import fitz  # PyMuPDF
    except Exception:
        return ""
    out: list[str] = []
    try:
        doc = fitz.open(path)
    except Exception:
        return ""
    for page in doc:
        txt = (page.get_text() or "").strip()
        if len(txt) < 40:  # likely a scanned page → OCR the rendered image
            try:
                from PIL import Image
                import io
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                txt = (_ocr_image(img) or txt).strip()
            except Exception:
                pass
        if txt:
            out.append(txt)
    doc.close()
    return "\n\n".join(out)


def _from_docx(path: Path) -> str:
    try:
        import docx
    except Exception:
        return ""
    try:
        d = docx.Document(str(path))
    except Exception:
        return ""
    parts = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    for tbl in d.tables:
        for row in tbl.rows:
            cells = [c.text.strip() for c in row.cells if c.text and c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_xlsx(path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, read_only=True, data_only=True)
    except Exception:
        return ""
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c).strip() for c in row if c not in (None, "")]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)


def _from_image(path: Path) -> str:
    try:
        from PIL import Image
        return _ocr_image(Image.open(path))
    except Exception:
        return ""


def _from_html(path: Path) -> str:
    raw = path.read_text("utf-8", errors="ignore")
    raw = re.sub(r"(?is)<(script|style)[^>]*>.*?</\1>", " ", raw)
    raw = re.sub(r"(?s)<[^>]+>", " ", raw)
    raw = re.sub(r"&nbsp;", " ", raw)
    return re.sub(r"[ \t]*\n\s*\n\s*", "\n\n", re.sub(r"[ \t]+", " ", raw)).strip()


def _extract_raw(path: Path) -> str:
    ext = path.suffix.lower()
    if ext in _TEXT_EXT:
        return path.read_text("utf-8", errors="ignore")
    if ext in _HTML_EXT:
        return _from_html(path)
    if ext == ".pdf":
        return _from_pdf(path)
    if ext == ".docx":
        return _from_docx(path)
    if ext in (".xlsx", ".xlsm"):
        return _from_xlsx(path)
    if ext in _IMG_EXT:
        return _from_image(path)
    return ""


def extract_text(path: Path) -> str:
    """Return the plain text of `path`, cached by (content hash + version)."""
    try:
        raw_bytes = path.read_bytes()
    except Exception:
        return ""
    h = hashlib.sha256(EXTRACT_VERSION.encode() + raw_bytes).hexdigest()[:16]
    cf = _CACHE_DIR / f"{h}.txt"
    if cf.exists():
        try:
            return cf.read_text("utf-8")
        except Exception:
            pass
    text = _extract_raw(path) or ""
    text = re.sub(r"\n{3,}", "\n\n", text).strip()
    try:
        _CACHE_DIR.mkdir(parents=True, exist_ok=True)
        cf.write_text(text, "utf-8")
    except Exception:
        pass
    return text


def is_supported(path: Path) -> bool:
    return path.suffix.lower() in (
        _TEXT_EXT | _HTML_EXT | _IMG_EXT | {".pdf", ".docx", ".xlsx", ".xlsm"})


def scan_docs() -> list[dict]:
    """Every supported document under the docs folder(s): {source, path, text}."""
    out: list[dict] = []
    for d in docs_dirs():
        for f in sorted(d.rglob("*")):
            if not f.is_file() or not is_supported(f):
                continue
            text = extract_text(f)
            if text and len(text) > 40:
                out.append({"source": f.name, "path": str(f), "text": text})
    return out


def status() -> dict:
    dirs = docs_dirs()
    files, extracted = 0, 0
    for d in dirs:
        for f in d.rglob("*"):
            if f.is_file() and is_supported(f):
                files += 1
    for doc in scan_docs():
        extracted += 1
    return {
        "dirs": [str(d) for d in dirs] or [str(DOCS_DIR)],
        "exists": bool(dirs),
        "files": files,
        "extracted": extracted,
        "tesseract": bool(_tesseract_exe()),
    }
